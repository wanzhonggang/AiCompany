import os
import logging
import re
from typing import Dict, List, Optional, Any
from pathlib import Path
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from .models import (
    KnowledgeBase,
    KnowledgeDocument,
    KnowledgeDocumentStatus,
    DocumentChunk,
)
from .time_utils import now_beijing
from .database import async_session
from .config import DATA_DIR

# Import document processing libraries
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

logger = logging.getLogger(__name__)

# Create upload directory under backend/data so runtime files stay out of source paths.
UPLOAD_DIR = DATA_DIR / "uploads" / "documents"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


async def create_knowledge_base(
    session: AsyncSession,
    enterprise_id: str,
    name: str,
    description: str = "",
    is_public: bool = False,
) -> KnowledgeBase:
    """Create a new knowledge base."""
    kb = KnowledgeBase(
        enterprise_id=enterprise_id,
        name=name,
        description=description,
        is_public=is_public,
    )
    session.add(kb)
    await session.commit()
    await session.refresh(kb)
    return kb


async def get_knowledge_bases(
    session: AsyncSession,
    enterprise_id: str,
) -> List[KnowledgeBase]:
    """Get all knowledge bases for an enterprise."""
    stmt = select(KnowledgeBase).where(KnowledgeBase.enterprise_id == enterprise_id)
    result = await session.execute(stmt)
    return list(result.scalars().all())


async def get_knowledge_base(
    session: AsyncSession,
    kb_id: str,
    enterprise_id: str,
) -> Optional[KnowledgeBase]:
    """Get a knowledge base by ID."""
    stmt = select(KnowledgeBase).where(
        KnowledgeBase.id == kb_id,
        KnowledgeBase.enterprise_id == enterprise_id,
    )
    result = await session.execute(stmt)
    return result.scalar_one_or_none()


async def add_document(
    session: AsyncSession,
    knowledge_base_id: str,
    filename: str,
    file_content: bytes,
    enterprise_id: str,
) -> KnowledgeDocument:
    """Add a document to a knowledge base."""
    # Validate knowledge base ownership
    kb = await get_knowledge_base(session, knowledge_base_id, enterprise_id)
    if not kb:
        raise ValueError("Knowledge base not found")
    
    # Save file
    safe_name = Path(filename).name.strip() or "upload"
    safe_name = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", safe_name)
    file_ext = Path(safe_name).suffix.lower()
    file_path = UPLOAD_DIR / f"{now_beijing().strftime('%Y%m%d%H%M%S')}_{safe_name}"
    
    with open(file_path, "wb") as f:
        f.write(file_content)
    
    # Create document record
    doc = KnowledgeDocument(
        knowledge_base_id=knowledge_base_id,
        filename=safe_name,
        file_path=str(file_path),
        file_type=file_ext,
        file_size=len(file_content),
        status=KnowledgeDocumentStatus.PENDING.value,
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)
    
    return doc


async def process_document_task(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Process a document task for the queue."""
    document_id = payload["document_id"]
    
    async with async_session() as session:
        stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        result = await session.execute(stmt)
        doc = result.scalar_one_or_none()
        
        if not doc:
            raise ValueError("Document not found")
        
        doc.status = KnowledgeDocumentStatus.PROCESSING.value
        await session.commit()
        
        try:
            # Extract text from document
            text = await _extract_text(doc.file_path, doc.file_type)
            doc.content = text
            
            # Split into chunks
            chunks = _split_text_into_chunks(text)
            
            # Save chunks
            for idx, chunk_text in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk_text,
                )
                session.add(chunk)
            
            doc.status = KnowledgeDocumentStatus.COMPLETED.value
            await session.commit()
            
            return {
                "document_id": doc.id,
                "chunks_count": len(chunks),
            }
            
        except Exception as e:
            logger.error(f"Failed to process document {document_id}: {str(e)}", exc_info=True)
            doc.status = KnowledgeDocumentStatus.FAILED.value
            doc.error_message = str(e)
            await session.commit()
            raise


async def _extract_text(file_path: str, file_type: str) -> str:
    """Extract text from document based on file type."""
    text = ""
    
    if file_type == ".pdf" and HAS_PDF:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
            
    elif file_type == ".docx" and HAS_DOCX:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
            
    elif file_type in [".xlsx", ".xls"] and HAS_XLSX:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"Sheet: {sheet_name}\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                text += row_text + "\n"
                
    elif file_type in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return text


def _split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = min(start + chunk_size, text_length)
        
        # Try to find a natural break (sentence end or space)
        if end < text_length:
            # Look for sentence ending within the last 100 characters
            last_period = text.rfind(".", start, end)
            last_space = text.rfind(" ", start, end)
            
            if last_period > start + chunk_size // 2:
                end = last_period + 1
            elif last_space > start + chunk_size // 2:
                end = last_space
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap if end < text_length else end
    
    return chunks


async def search_knowledge(
    session: AsyncSession,
    enterprise_id: str,
    query: str,
    knowledge_base_ids: Optional[List[str]] = None,
    top_k: int = 5,
) -> List[Dict[str, Any]]:
    """Simple keyword-based search for knowledge (placeholder for vector search)."""
    # Build query
    stmt = select(DocumentChunk).join(KnowledgeDocument).join(KnowledgeBase).where(
        KnowledgeBase.enterprise_id == enterprise_id,
    )
    
    if knowledge_base_ids:
        stmt = stmt.where(KnowledgeBase.id.in_(knowledge_base_ids))
    
    # Simple keyword matching
    keywords = [kw.lower() for kw in query.split() if kw.strip()]
    
    result = await session.execute(stmt.options(
        selectinload(DocumentChunk.document).selectinload(KnowledgeDocument.knowledge_base)
    ))
    chunks = result.scalars().all()
    
    # Score chunks based on keyword matches
    scored_chunks = []
    for chunk in chunks:
        score = 0.0
        chunk_text = chunk.content.lower()
        for keyword in keywords:
            if keyword in chunk_text:
                score += 1.0
        if score > 0:
            scored_chunks.append({
                "chunk": chunk,
                "document": chunk.document,
                "score": score,
            })
    
    # Sort and limit
    scored_chunks.sort(key=lambda x: x["score"], reverse=True)
    return scored_chunks[:top_k]


async def get_documents(
    session: AsyncSession,
    knowledge_base_id: str,
    enterprise_id: str,
) -> List[KnowledgeDocument]:
    """Get all documents in a knowledge base."""
    kb = await get_knowledge_base(session, knowledge_base_id, enterprise_id)
    if not kb:
        raise ValueError("Knowledge base not found")
    
    stmt = select(KnowledgeDocument).where(KnowledgeDocument.knowledge_base_id == knowledge_base_id)
    result = await session.execute(stmt)
    return list(result.scalars().all())
